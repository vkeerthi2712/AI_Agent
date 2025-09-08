import os
import io
import datetime
import random
import time
import pandas as pd
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException
from fastapi.responses import StreamingResponse
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Emu
from docx.document import Document as DocxDocument
from base64 import b64decode, b64encode

# Optional PDF conversion
try:
    from docx2pdf import convert
except ImportError:
    convert = None

import requests

app = FastAPI(title="AI-Powered Report Generator")

# Load environment variables from .env file
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
GEN_API_BASE = "https://generativelanguage.googleapis.com/v1"

if not GEMINI_API_KEY:
    raise ValueError("❌ GEMINI_API_KEY not found in .env file.")

# --- Gemini API Helper Functions ---
def _gemini_generate_content(api_key: str, model_id: str, prompt: str) -> dict:
    """Sends a request to the Gemini API with retry logic."""
    url = f"{GEN_API_BASE}/models/{model_id}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}
    max_attempts = 5
    for attempt in range(max_attempts):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=60)
            if resp.status_code in (429, 503):
                time.sleep((2 ** attempt) + random.random())
                continue
            resp.raise_for_status()
            return resp.json()
        except requests.exceptions.RequestException as e:
            if attempt == max_attempts - 1:
                raise HTTPException(status_code=500, detail=f"Gemini API error: {e}")
            time.sleep((2 ** attempt) + random.random())
    raise HTTPException(status_code=500, detail="Gemini API retries exhausted.")

def get_gemini_text(api_key: str, prompt: str, model_id: str) -> str:
    """Extracts text content from a Gemini API response."""
    result = _gemini_generate_content(api_key, model_id, prompt)
    try:
        return result["candidates"][0]["content"]["parts"][0]["text"]
    except (KeyError, IndexError):
        return "No summary available."

# --- Report Generation Logic ---
def create_report_document(df: pd.DataFrame, user_query: str, model_id: str, chart_image_base64: str = None) -> DocxDocument:
    """Generates a Word document with data and AI-powered insights, including an embedded chart."""
    doc = Document()
    df_string = df.to_csv(index=False)

    # Use Gemini to generate the executive summary and conclusion
    summary_prompt = f"Write a professional executive summary for user's request: '{user_query}' with the following data insights:\n{df_string}"
    summary_text = get_gemini_text(GEMINI_API_KEY, summary_prompt, model_id)

    conclusion_prompt = f"Write a professional and concise final recommendation or summary statement for user's request: '{user_query}' with the following data:\n{df_string}"
    conclusion_text = get_gemini_text(GEMINI_API_KEY, conclusion_prompt, model_id)

    # --- Cover Page ---
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("Customer Name + Logo Placeholder")
    run_logo.font.size = Pt(20)
    run_logo.font.bold = True
    doc.add_heading("Report Title", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared by Innov® Labs AI Agent").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # --- Executive Summary ---
    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary_text)

    # --- Data Insights ---
    doc.add_page_break()
    doc.add_heading("Data Insights", level=1)
    doc.add_paragraph("Tables, key statistics, interpretations")

    # Data Table
    doc.add_paragraph("Data Table:")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val).replace('\n', ' ')

    # Key Statistics
    doc.add_paragraph("\nKey Statistics:")
    stats = df.describe(include='all').fillna('')
    stats_table = doc.add_table(rows=1, cols=len(stats.columns) + 1)
    stats_table.style = 'Table Grid'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = "Statistic"
    for i, col in enumerate(stats.columns):
        hdr_cells[i + 1].text = str(col)
    for stat_name, row in stats.iterrows():
        row_cells = stats_table.add_row().cells
        row_cells[0].text = str(stat_name)
        for i, val in enumerate(row):
            row_cells[i + 1].text = str(round(val, 2)) if pd.api.types.is_numeric_dtype(val) else str(val)

    # Interpretation
    doc.add_paragraph("\nInterpretation:")
    doc.add_paragraph("The table above shows descriptive statistics of the dataset, including mean, std, min, max, and quartiles.")

    # --- Charts / Visualizations ---
    doc.add_page_break()
    doc.add_heading("Charts / Visualizations", level=1)
    if chart_image_base64:
        # Decode the base64 string and add the image to the document
        try:
            image_bytes = b64decode(chart_image_base64)
            image_stream = io.BytesIO(image_bytes)
            # Add image with a fixed size to maintain aspect ratio
            doc.add_picture(image_stream, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Error embedding chart image: {e}")
    else:
        doc.add_paragraph("No chart image was provided for the report.")
        
    doc.add_paragraph("AI-generated plots for key insights from the data.")

    # --- Conclusion ---
    doc.add_page_break()
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(conclusion_text)

    # --- Verification / Audit ---
    doc.add_page_break()
    doc.add_heading("Verification / Audit", level=1)
    doc.add_paragraph("Prepared by: Innov® Labs AI Agent")
    doc.add_paragraph("Verified by: _______________________ (Signature)")
    doc.add_paragraph(f"Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    return doc

# --- FastAPI Endpoints ---
@app.get("/health")
def health():
    """Endpoint to check the server status."""
    return {"status": "ok", "model": GEMINI_MODEL}

@app.post("/get_insights")
async def get_insights_endpoint(
    file: UploadFile = File(...),
    user_query: str = Form(...),
    model_id: str = Form(None)
):
    """Generates and returns AI-powered insights and a chart specification."""
    if not file.filename.endswith(('.csv', '.xlsx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Upload CSV or XLSX.")

    chosen_model = model_id or GEMINI_MODEL

    try:
        raw_data = await file.read()
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.StringIO(raw_data.decode('utf-8')))
        else:
            df = pd.read_excel(io.BytesIO(raw_data))

        # Generate insights text
        insights_prompt = f"Provide a concise, single-paragraph analysis of the following data based on the user's query:\nQuery: '{user_query}'\nData:\n{df.to_csv(index=False)}"
        insights_text = get_gemini_text(GEMINI_API_KEY, insights_prompt, chosen_model)
        
        # Generate chart specification
        chart_prompt = f"Based on the following data and user query, provide a JSON object for an Altair chart. Use a 'bar' or 'line' chart type. Select the two most relevant columns for a simple X/Y plot. The object should have a 'chart_type', a 'title', an 'x_axis', and a 'y_axis'.  Use column names from the data. Example: {{'chart_type': 'bar', 'title': 'Sales by Product', 'x_axis': 'Product', 'y_axis': 'Sales'}} \nQuery: '{user_query}'\nData Columns: {list(df.columns)}"
        chart_spec_json = get_gemini_text(GEMINI_API_KEY, chart_prompt, chosen_model)

        return {
            "insights": insights_text,
            "chart_spec": chart_spec_json,
            "data_preview": df.to_json(orient='records')
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.post("/generate_report")
async def generate_report_endpoint(
    file: UploadFile = File(...),
    user_query: str = Form(...),
    report_format: str = Form("docx"),
    model_id: str = Form(None),
    # CHANGE: Accept chart_image as an UploadFile object, which is a bytes-like object.
    chart_image: UploadFile = File(None) 
):
    """Generates an AI-powered report from uploaded data."""
    if not file.filename.endswith(('.csv', '.xlsx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Upload CSV or XLSX.")

    chosen_model = model_id or GEMINI_MODEL
    chart_image_base64 = None

    try:
        raw_data = await file.read()
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.StringIO(raw_data.decode('utf-8')))
        else:
            df = pd.read_excel(io.BytesIO(raw_data))

        # CHANGE: If a chart image was uploaded, read it and convert it to base64.
        if chart_image:
            image_bytes = await chart_image.read()
            chart_image_base64 = b64encode(image_bytes).decode('utf-8')

        # Pass the base64 string to the document creation function
        doc = create_report_document(df, user_query, chosen_model, chart_image_base64)

        if report_format.lower() == "pdf":
            if convert is None:
                raise HTTPException(status_code=500, detail="PDF generation requires 'docx2pdf' and Word installed.")
            temp_docx = "temp_report.docx"
            temp_pdf = "temp_report.pdf"
            doc.save(temp_docx)
            convert(temp_docx, temp_pdf)
            with open(temp_pdf, "rb") as f:
                file_bytes = f.read()
            filename = "generated_report.pdf"
        else:
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            file_bytes = doc_io.getvalue()
            filename = "generated_report.docx"

        return StreamingResponse(
            io.BytesIO(file_bytes),
            media_type="application/pdf" if report_format.lower() == "pdf" else "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"}
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")
