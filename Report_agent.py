# report_agent.py
import os
import io
import datetime
import random
import time
import pandas as pd
import json
from dotenv import load_dotenv
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Request
from fastapi.responses import StreamingResponse, JSONResponse
from pydantic import BaseModel
from typing import Optional
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from base64 import b64decode, b64encode

try:
    from docx2pdf import convert
except ImportError:
    convert = None

import requests
import traceback

app = FastAPI(title="AI-Powered Report Generator with HITL")

# Load environment variables
load_dotenv()
GEMINI_API_KEY = os.getenv("GEMINI_API_KEY")
GEMINI_MODEL = os.getenv("GEMINI_MODEL", "gemini-1.5-flash")
GEN_API_BASE = "https://generativelanguage.googleapis.com/v1"

if not GEMINI_API_KEY:
    # For local testing you might prefer not to raise, but warn
    print("Warning: GEMINI_API_KEY not found. Gemini calls will fail.")
    # raise ValueError("❌ GEMINI_API_KEY not found in .env file.")

# --- In-memory HITL store (simple) ---
# Structure:
# pending_reviews = {
#   review_id: {
#       "insights": str,
#       "chart_spec": dict,
#       "data_preview": <json string>,
#       "approved": False,
#       "reviewer_edits": None,
#       "created_at": iso,
#       "reviewed_at": iso or None,
#       "reviewed_by": str or None
#   }
# }
pending_reviews = {}

# --- Gemini helpers (robust with fallback) ---
def _gemini_generate_content(api_key: str, model_id: str, prompt: str) -> dict:
    """Call Gemini. On repeated failures, return object with hitl_required flag."""
    if not api_key:
        return {"hitl_required": True, "error": "Missing API key", "prompt": prompt}

    url = f"{GEN_API_BASE}/models/{model_id}:generateContent?key={api_key}"
    headers = {"Content-Type": "application/json"}
    payload = {"contents": [{"role": "user", "parts": [{"text": prompt}]}]}
    max_attempts = 4
    for attempt in range(max_attempts):
        try:
            resp = requests.post(url, headers=headers, json=payload, timeout=30)
            if resp.status_code in (429, 503):
                # exponential backoff
                time.sleep((2 ** attempt) + random.random())
                continue
            resp.raise_for_status()
            return resp.json()
        except requests.exceptions.RequestException as e:
            if attempt == max_attempts - 1:
                return {"hitl_required": True, "error": str(e), "prompt": prompt}
            time.sleep((2 ** attempt) + random.random())
    return {"hitl_required": True, "error": "Unknown failure", "prompt": prompt}

def get_gemini_text(api_key: str, prompt: str, model_id: str) -> dict:
    """Return dict: {ok:bool, text:..., hitl_required:bool, error:...}"""
    res = _gemini_generate_content(api_key, model_id, prompt)
    if res.get("hitl_required"):
        return {"ok": False, "text": None, "hitl_required": True, "error": res.get("error"), "prompt": res.get("prompt")}
    try:
        text = res["candidates"][0]["content"]["parts"][0]["text"]
        return {"ok": True, "text": text, "hitl_required": False}
    except Exception as e:
        return {"ok": False, "text": None, "hitl_required": True, "error": f"Parse error: {e}"}

def get_gemini_json(api_key: str, prompt: str, model_id: str) -> dict:
    """Return dict with chart_spec or a hitl_required flag."""
    res = _gemini_generate_content(api_key, model_id, prompt)
    if res.get("hitl_required"):
        return {"hitl_required": True, "error": res.get("error"), "prompt": res.get("prompt")}
    try:
        raw_text = res["candidates"][0]["content"]["parts"][0]["text"]
        # strip code fences
        json_string = raw_text.strip().lstrip('```json').rstrip('```').strip()
        return {"hitl_required": False, "chart_spec": json.loads(json_string)}
    except Exception as e:
        return {"hitl_required": True, "error": f"Failed to parse JSON: {e}", "raw_text": raw_text if 'raw_text' in locals() else None}

# --- DOCX helpers ---
def add_watermark_to_doc(doc: Document, watermark_text: str = "Innov® Labs"):
    for section in doc.sections:
        header = section.header
        paragraph = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        paragraph.text = watermark_text
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = Pt(40)
        run.font.color.rgb = RGBColor(200, 200, 200)
        run.font.bold = True
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        hdr = header._element
        p = paragraph._element
        pPr = OxmlElement('w:pPr')
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.insert(0, pPr)

def create_report_document(df: pd.DataFrame, user_query: str, model_id: str,
                           chart_image_base64: str = None,
                           summary_text: str = None,
                           conclusion_text: str = None) -> Document:
    doc = Document()
    data_to_send = df.copy()
    if model_id == "gemini-1.5-pro":
        data_to_send = df.head(50)
    df_string = data_to_send.to_csv(index=False)

    if not summary_text:
        s = get_gemini_text(GEMINI_API_KEY, f"Write an executive summary for: '{user_query}'\nData:\n{df_string}", model_id)
        summary_text = s.get("text") or "No AI summary available."

    if not conclusion_text:
        c = get_gemini_text(GEMINI_API_KEY, f"Write a concise conclusion for: '{user_query}'\nData:\n{df_string}", model_id)
        conclusion_text = c.get("text") or "No AI conclusion available."

    # Cover
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_logo = p_logo.add_run("Customer Name")
    run_logo.font.size = Pt(20)
    run_logo.font.bold = True
    doc.add_heading("Report Title", level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Date: {datetime.date.today().strftime('%B %d, %Y')}").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Prepared by Innov® Labs AI Agent").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    doc.add_heading("Executive Summary", level=1)
    doc.add_paragraph(summary_text)

    doc.add_page_break()
    doc.add_heading("Data Insights", level=1)
    doc.add_paragraph("\nData Table:")
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = 'Table Grid'
    hdr_cells = table.rows[0].cells
    for i, col in enumerate(df.columns):
        hdr_cells[i].text = str(col)
    for _, row in df.iterrows():
        row_cells = table.add_row().cells
        for i, val in enumerate(row):
            row_cells[i].text = str(val).replace('\n', ' ')
    doc.add_paragraph("\nKey Statistics:")
    stats = df.describe(include='all').fillna('')
    stats_table = doc.add_table(rows=1, cols=len(stats.columns) + 1)
    stats_table.style = 'Table Grid'
    hdr_cells = stats_table.rows[0].cells
    hdr_cells[0].text = "Statistic"
    for i, col in enumerate(stats.columns):
        hdr_cells[i + 1].text = str(col)
    for stat_name, row in stats.iterrows():
        row_cells = stats_table.add_row().cells
        row_cells[0].text = str(stat_name)
        for i, val in enumerate(row):
            row_cells[i + 1].text = str(round(val, 2)) if pd.api.types.is_numeric_dtype(val) else str(val)

    doc.add_paragraph("\nInterpretation:")
    doc.add_paragraph("The tables above show descriptive statistics and a raw data preview of the dataset.")

    doc.add_page_break()
    doc.add_heading("Charts / Visualizations", level=1)
    if chart_image_base64:
        try:
            image_bytes = b64decode(chart_image_base64)
            image_stream = io.BytesIO(image_bytes)
            doc.add_picture(image_stream, width=Inches(6))
        except Exception as e:
            doc.add_paragraph(f"Error embedding chart image: {e}")
    else:
        doc.add_paragraph("No chart image provided.")

    doc.add_paragraph("AI-generated plots for key insights from the data.")

    doc.add_page_break()
    doc.add_heading("Conclusion", level=1)
    doc.add_paragraph(conclusion_text)

    doc.add_page_break()
    doc.add_heading("Verification / Audit", level=1)
    doc.add_paragraph("Prepared by: Innov® Labs AI Agent")
    doc.add_paragraph("Verified by: _______________________ (Signature)")
    doc.add_paragraph(f"Timestamp: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")

    add_watermark_to_doc(doc, "Innov® Labs")
    return doc

# --- Models for JSON review submission ---
class ReviewSubmission(BaseModel):
    review_id: str
    approved: bool
    reviewer_name: Optional[str] = "anonymous"
    # reviewer_edits can contain {"summary": "...", "conclusion": "...", "chart_spec": {...}}
    reviewer_edits: Optional[dict] = None

# --- Endpoints ---
@app.get("/health")
def health():
    return {"status": "ok", "model": GEMINI_MODEL}

@app.post("/get_insights")
async def get_insights_endpoint(
    file: UploadFile = File(...),
    user_query: str = Form(...),
    model_id: Optional[str] = Form(None)
):
    """Generate a draft (AI) and add to pending_reviews for HITL."""
    if not file.filename.endswith(('.csv', '.xlsx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Upload CSV or XLSX.")
    chosen_model = model_id or GEMINI_MODEL
    try:
        raw_data = await file.read()
        # parse file contents into df
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.StringIO(raw_data.decode('utf-8')))
        else:
            df = pd.read_excel(io.BytesIO(raw_data))

        # Ask the LLM for a one-paragraph insight (wrapped with hitl fallback)
        insights_prompt = f"Provide a concise, single-paragraph analysis of this dataset for the user's query.\nQuery: {user_query}\nData (first 50 rows):\n{df.head(50).to_csv(index=False)}"
        insight_res = get_gemini_text(GEMINI_API_KEY, insights_prompt, chosen_model)
        if insight_res.get("hitl_required"):
            insights_text = None
            insight_error = insight_res.get("error", "AI unavailable")
        else:
            insights_text = insight_res.get("text")
            insight_error = None

        # Ask the LLM for chart spec JSON
        chart_prompt = f"Return JSON with fields chart_type ('bar' or 'line'), title, x_axis, y_axis. Query: {user_query}. Columns: {list(df.columns)}"
        chart_res = get_gemini_json(GEMINI_API_KEY, chart_prompt, chosen_model)
        if chart_res.get("hitl_required"):
            chart_spec = {"chart_type": "bar", "x_axis": df.columns[0] if len(df.columns)>0 else "", "y_axis": df.columns[1] if len(df.columns)>1 else ""}
            chart_error = chart_res.get("error")
            hitl_needed = True
        else:
            chart_spec = chart_res.get("chart_spec", {})
            chart_error = None
            hitl_needed = False

        review_id = str(random.randint(100000, 999999))
        pending_reviews[review_id] = {
            "insights": insights_text or "",
            "insight_error": insight_error,
            "chart_spec": chart_spec,
            "chart_error": chart_error,
            "data_preview": df.head(200).to_json(orient='records'),
            "approved": False,
            "reviewer_edits": None,
            "created_at": datetime.datetime.now().isoformat(),
            "reviewed_at": None,
            "reviewed_by": None
        }

        return JSONResponse(status_code=200, content={
            "review_id": review_id,
            "status": "pending_review",
            "draft": pending_reviews[review_id],
            "hitl_needed": hitl_needed or (insight_res.get("hitl_required") if isinstance(insight_res, dict) else False)
        })
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")

@app.post("/review_insights")
async def review_insights(payload: ReviewSubmission):
    """Endpoint to accept human reviewer edits (JSON) or approval flag."""
    review_id = payload.review_id
    if review_id not in pending_reviews:
        raise HTTPException(status_code=404, detail="Review ID not found")
    record = pending_reviews[review_id]
    record["approved"] = bool(payload.approved)
    record["reviewed_at"] = datetime.datetime.now().isoformat()
    record["reviewed_by"] = payload.reviewer_name
    # store reviewer_edits as JSON dict
    record["reviewer_edits"] = payload.reviewer_edits or None
    return {"review_id": review_id, "status": "approved" if payload.approved else "rejected", "record": record}

@app.post("/generate_report")
async def generate_report_endpoint(
    file: UploadFile = File(...),
    user_query: str = Form(...),
    report_format: str = Form("docx"),
    model_id: Optional[str] = Form(None),
    chart_image: UploadFile = File(None),
    review_id: Optional[str] = Form(None)
):
    """Generate a report. If review_id present and approved, use reviewer edits (summary/conclusion/chart)"""
    if not file.filename.endswith(('.csv', '.xlsx')):
        raise HTTPException(status_code=400, detail="Invalid file type. Upload CSV or XLSX.")
    chosen_model = model_id or GEMINI_MODEL
    chart_image_base64 = None
    temp_docx = None
    temp_pdf = None
    try:
        file_bytes = await file.read()
        file_stream = io.BytesIO(file_bytes)
        if file.filename.endswith('.csv'):
            df = pd.read_csv(io.StringIO(file_bytes.decode('utf-8')))
        else:
            df = pd.read_excel(file_stream)

        if chart_image:
            image_bytes = await chart_image.read()
            chart_image_base64 = b64encode(image_bytes).decode('utf-8')

        summary_text = None
        conclusion_text = None
        # If review_id provided and approved, take edits
        if review_id:
            record = pending_reviews.get(review_id)
            if record and record.get("approved"):
                edits = record.get("reviewer_edits")
                if isinstance(edits, dict):
                    summary_text = edits.get("summary")
                    conclusion_text = edits.get("conclusion")
                    # If reviewer provided chart_spec, use that to embed? (we still accept chart_image upload)
                # else ignore malformed
        # Create doc using optional human edits
        doc = create_report_document(df, user_query, chosen_model, chart_image_base64, summary_text, conclusion_text)

        if report_format.lower() == "pdf":
            if convert is None:
                raise HTTPException(status_code=500, detail="PDF generation requires docx2pdf + local Word installation.")
            temp_docx = f"temp_{random.randint(100000, 999999)}.docx"
            temp_pdf = f"temp_{random.randint(100000, 999999)}.pdf"
            doc.save(temp_docx)
            convert(temp_docx, temp_pdf)
            with open(temp_pdf, "rb") as f:
                file_bytes = f.read()
            filename = "generated_report.pdf"
            media_type = "application/pdf"
        else:
            doc_io = io.BytesIO()
            doc.save(doc_io)
            doc_io.seek(0)
            file_bytes = doc_io.getvalue()
            filename = "generated_report.docx"
            media_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

        return StreamingResponse(io.BytesIO(file_bytes), media_type=media_type, headers={"Content-Disposition": f"attachment; filename={filename}"})
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"An error occurred: {str(e)}")
    finally:
        if temp_docx and os.path.exists(temp_docx):
            os.remove(temp_docx)
        if temp_pdf and os.path.exists(temp_pdf):
            os.remove(temp_pdf)
